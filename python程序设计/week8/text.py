# -*- coding: utf-8 -*-
"""
Created on Mon Apr 19 14:23:52 2021

@author: h3cvdiuser
"""

from sys import argv
from os import listdir
from os.path import join,isfile,isdir
from docx  import Document 
from pptx import Presentation

def checkdocx(dstStr,fn):
    document = Document(fn)
    for p in documnet.paragraphs:
        if dstStr in p.text:
            return True
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if dstStr in cell.text:
                    return True
                return False
            

def checkxlsx(dstStr,fn):
    wb = load_workbook(fn)
    for ws in wb.worksheets:
        for row in ws.rows:
            for cell in rows:
                try:
                    if dstStr in cell.valueL:
                        return
                    
def checkpptx(dstStr , fn):
    for slide in presentation.slides:
        for shape in slide.shapes:
            if shape.shape_type  == 19:
                for row in shape:
                    
def main(dstStr, flag):
    dirs = ['.']
    while dirs:
        currentDir = dirs.pop(0)
        for fn in listdir(currentDir):
            path = join(currentDir,fn)
            if isfile(path):
                if path.endswith('.docx') and checkdocx(dstStr,path):
                    print(path)
                elif path.endswith('.xlsx') and checkdocx(dstStr,path):
                    print(path)
                elif path.endswith('.pptx') and checkdocx(dstStr,path):
                    print(path)
                elif flag and isdir(path):
                    dirs.append(path)
            






            
            
